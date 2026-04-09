"""
Digest DOCX renderer.

Takes a DigestResult and writes a formatted .docx file that matches
the official "Mau Tong thuat Book" template used by Học viện KTQS.

Layout
------
  Header block  — HỌC VIỆN KỸ THUẬT QUÂN SỰ / PHÒNG THÔNG TIN …
  ─────────────────────────────────────────────────────────────────
  1. THÔNG TIN CHUNG VỀ TÀI LIỆU
       - Nhan đề / Tác giả / NXB / … (from Document metadata)
  2. TỔNG THUẬT VỀ TÀI LIỆU
     2.1 Tóm tắt
     2.2 Nội dung chính
     2.3 Từ khóa
  3. PHẠM VI SỬ DỤNG
       - Hướng nghiên cứu  (from research_directions)
  * Thông tin quản trị CSDL

Dependency:  python-docx (already in requirements.txt)
"""
import io
from typing import Optional

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from services.digest_service import DigestResult


# ── Helpers ───────────────────────────────────────────────────────────

def _bold(paragraph, text: str):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def _add_heading(doc: DocxDocument, text: str, level: int = 1):
    """Add a numbered section heading (bold, 13 pt)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 12)
    return p


def _add_subheading(doc: DocxDocument, text: str):
    """Sub-heading like '2.1. Tóm tắt' (bold, 12 pt, slight indent)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def _add_bullet(doc: DocxDocument, text: str, indent: int = 1):
    """Bullet line (- …) with configurable indent level."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(12 * indent)
    p.add_run(text)
    return p


def _add_field(doc: DocxDocument, label: str, value: Optional[str]):
    """  - Label: value"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    run = p.add_run(f"- {label}: ")
    run.bold = True
    p.add_run(value or "")
    return p


def _add_body(doc: DocxDocument, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text or "")
    return p


def _add_separator(doc: DocxDocument):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("*" * 5)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


# ── Renderer ──────────────────────────────────────────────────────────

class DigestRenderer:
    """
    Render a DigestResult to a .docx byte stream.

    Usage:
        renderer = DigestRenderer()
        docx_bytes = renderer.render(digest)
        with open("digest.docx", "wb") as f:
            f.write(docx_bytes)
    """

    def render(self, digest: DigestResult, reviewer: str = "", reviewer_approved: str = "") -> bytes:
        """
        Build a .docx in memory and return the raw bytes.

        Parameters
        ----------
        digest          DigestResult from DigestService.assemble()
        reviewer        Name of the person who reviewed / edited (for admin block)
        reviewer_approved  Name of the approver (for admin block)
        """
        doc = DocxDocument()

        # ── Default style: Times New Roman 12pt ─────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        # ── Header block ─────────────────────────────────────────────
        h1 = doc.add_paragraph()
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h1.add_run("HỌC VIỆN KỸ THUẬT QUÂN SỰ")
        r.bold = True
        r.font.size = Pt(13)

        h2 = doc.add_paragraph()
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h2.add_run("PHÒNG THÔNG TIN KHOA HỌC QUÂN SỰ")

        _add_separator(doc)

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = title_p.add_run("TỔNG THUẬT TÀI LIỆU")
        r2.bold = True
        r2.font.size = Pt(14)

        doc.add_paragraph()  # blank line

        # ══════════════════════════════════════════════════════════════
        # 1. THÔNG TIN CHUNG VỀ TÀI LIỆU
        # ══════════════════════════════════════════════════════════════
        _add_heading(doc, "1. THÔNG TIN CHUNG VỀ TÀI LIỆU", level=1)

        _add_field(doc, "Nhan đề (Title)", digest.title)
        _add_field(doc, "Tác giả (Authors)", "")
        _add_field(doc, "Nhà xuất bản (Publisher)", "")
        _add_field(doc, "Năm xuất bản (Year)", "")
        _add_field(doc, "ISBN", "")
        _add_field(doc, "DOI", "")
        _add_field(doc, "Số trang (Pages)", "")
        _add_field(doc, "Ngôn ngữ gốc (Language)", digest.source_language.upper())

        doc.add_paragraph()

        # ══════════════════════════════════════════════════════════════
        # 2. TỔNG THUẬT VỀ TÀI LIỆU
        # ══════════════════════════════════════════════════════════════
        _add_heading(doc, "2. TỔNG THUẬT VỀ TÀI LIỆU", level=1)

        # ── 2.1 Abstract ─────────────────────────────────────────────
        _add_subheading(doc, "2.1. Tóm tắt")
        if digest.abstract:
            _add_body(doc, digest.abstract)
        else:
            _add_body(doc, "[Chưa có — chạy summarize trước]")

        doc.add_paragraph()

        # ── 2.2 Main content breakdown ───────────────────────────────
        _add_subheading(doc, "2.2. Nội dung chính của tài liệu")

        if digest.main_content:
            mc = digest.main_content

            key_points = mc.get("key_points") or []
            if key_points:
                _add_body(doc, "Các điểm chính:")
                for kp in key_points:
                    _add_bullet(doc, str(kp), indent=2)

            methods = mc.get("methods") or []
            if methods:
                _add_body(doc, "Phương pháp / cách tiếp cận:")
                for m in methods:
                    _add_bullet(doc, str(m), indent=2)

            results = mc.get("results") or []
            if results:
                _add_body(doc, "Kết quả / phát hiện:")
                for r in results:
                    _add_bullet(doc, str(r), indent=2)

            conclusions = mc.get("conclusions") or []
            if conclusions:
                _add_body(doc, "Kết luận:")
                for c in conclusions:
                    _add_bullet(doc, str(c), indent=2)
        else:
            _add_body(doc, "[Chưa có — chạy main_content service trước]")

        doc.add_paragraph()

        # ── 2.3 Keywords ─────────────────────────────────────────────
        _add_subheading(doc, "2.3. Từ khóa")

        if digest.keywords:
            for entry in digest.keywords:
                _add_bullet(doc, entry.keyword, indent=2)
        else:
            _add_body(doc, "[Chưa có — chạy keyword service trước]")

        doc.add_paragraph()

        # ══════════════════════════════════════════════════════════════
        # 3. PHẠM VI SỬ DỤNG
        # ══════════════════════════════════════════════════════════════
        _add_heading(doc, "3. PHẠM VI SỬ DỤNG", level=1)

        _add_field(doc, "CTĐT đại học", "")
        _add_field(doc, "CTĐT thạc sĩ", "")
        _add_field(doc, "CTĐT tiến sĩ", "")
        _add_field(doc, "Nhóm nghiên cứu mạnh", "")

        # Research directions
        p_rd = doc.add_paragraph()
        p_rd.paragraph_format.left_indent = Pt(12)
        _bold(p_rd, "- Hướng nghiên cứu: ")

        if digest.research_directions:
            for entry in digest.research_directions:
                _add_bullet(doc, entry.direction_name, indent=2)
        else:
            _add_body(doc, "[Chưa có — chạy research_direction service trước]")

        doc.add_paragraph()

        # ══════════════════════════════════════════════════════════════
        # * Thông tin quản trị CSDL
        # ══════════════════════════════════════════════════════════════
        _add_separator(doc)

        admin_p = doc.add_paragraph()
        r_admin = admin_p.add_run("* Thông tin quản trị CSDL:")
        r_admin.bold = True

        _add_field(doc, "Người tổng thuật", "")
        _add_field(doc, "Người hiệu đính, phê duyệt", reviewer_approved)
        _add_field(doc, "Ngày nhập CSDL Học liệu số", "")

        # ── Missing sections warning ─────────────────────────────────
        if digest.missing:
            doc.add_paragraph()
            warn_p = doc.add_paragraph()
            r_w = warn_p.add_run("⚠ Các phần chưa xử lý:")
            r_w.bold = True
            r_w.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            for m in digest.missing:
                _add_bullet(doc, m, indent=1)

        # ── Serialize to bytes ───────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
