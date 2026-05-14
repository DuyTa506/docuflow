"""
Document management endpoints (v2).

POST   /api/v2/documents/upload
POST   /api/v2/documents/{id}/extract
POST   /api/v2/documents/{id}/text/upload   — Override OCR/extracted text via file
GET    /api/v2/documents
GET    /api/v2/documents/{id}
GET    /api/v2/documents/{id}/text
GET    /api/v2/documents/{id}/text/download — Download OCR/normalized text as .txt file
GET    /api/v2/documents/{id}/pages
GET    /api/v2/documents/{id}/elements
DELETE /api/v2/documents/{id}
"""
import os
import shutil
from typing import List, Optional

from fastapi import APIRouter, Depends, File, UploadFile, HTTPException, Query, Form
from fastapi import status as http_status
from fastapi.responses import Response
from sqlalchemy.orm import Session

from api.dependencies import get_db, get_current_user
from api.schemas import (
    TaskSubmittedResponse,
    DocumentUploadResponse,
    DocumentDetailResponse,
    DocumentTextResponse,
    DocumentListItem,
    PageListItem,
    ElementListItem,
)
from config.settings import settings
from data.db_models import User, Task
from data.repositories import DocumentRepository
from services.document_service import DocumentService
from utils.file_upload import extract_text_from_upload

router = APIRouter(prefix="/api/v2/documents", tags=["documents"])
_doc_svc = DocumentService()


# ── Upload ──────────────────────────────────────────────────────────

@router.post("/upload", response_model=DocumentUploadResponse, status_code=201)
async def upload_document(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    source_language: str = Form("en"),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Upload a PDF / image / DOCX / DOC file and register it as a document."""
    # Ensure upload directory exists
    os.makedirs(settings.upload_dir, exist_ok=True)

    # Validate file extension
    ext = os.path.splitext(file.filename)[1].lower()
    allowed = {".pdf", ".png", ".jpg", ".jpeg", ".docx", ".doc"}
    if ext not in allowed:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {sorted(allowed)}",
        )

    # Save file
    safe_name = file.filename.replace(" ", "_")
    dest = os.path.join(settings.upload_dir, safe_name)
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    doc = _doc_svc.upload_document(
        db,
        file_path_on_disk=dest,
        original_filename=file.filename,
        user_id=user.id,
        title=title,
        source_language=source_language,
    )

    return DocumentUploadResponse(
        document_id=doc.id,
        title=doc.title,
        format=doc.format,
        total_pages=doc.total_pages,
        processing_status=doc.processing_status,
    )


# ── Trigger unified extraction ──────────────────────────────────────

@router.post("/{document_id}/extract", response_model=TaskSubmittedResponse)
async def start_extraction(
    document_id: str,
    db: Session = Depends(get_db),
    _user: User = Depends(get_current_user),
):
    """
    Start unified extraction as a background task.

    Routes automatically based on document format:
    - DOCX / DOC  → python-docx (DOC first converted via LibreOffice)
    - PDF text    → PyMuPDF direct text extraction (per text page)
    - PDF scanned → DeepSeek vLLM OCR (per scanned page)
    - Image       → DeepSeek vLLM OCR
    """
    try:
        task_id = _doc_svc.submit_extraction(db, document_id)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc))
    return TaskSubmittedResponse(task_id=task_id, message="Extraction task submitted")


# ── List documents ──────────────────────────────────────────────────

@router.get("", response_model=list[DocumentListItem])
async def list_documents(
    limit: int = Query(50, ge=1, le=100),
    offset: int = Query(0, ge=0),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    List documents (paginated).

    Visibility rules:
    - ADMIN  → sees all documents from all users
    - MEMBER → sees only their own documents

    Each item includes a `task_summary` dict with the latest status of each
    pipeline task (EXTRACT, TRANSLATE, SUMMARIZE, KEYWORDS, etc.).
    """
    repo = DocumentRepository(db)

    if user.role == "ADMIN":
        docs = repo.list(limit=limit, offset=offset)
    else:
        docs = repo.list_for_user(user.id, limit=limit, offset=offset)

    # Build task summary for each document in one batch query
    doc_ids = [d.id for d in docs]
    task_summary_map: dict[str, dict[str, str]] = {doc_id: {} for doc_id in doc_ids}
    if doc_ids:
        # Fetch all tasks for these documents, newest first
        tasks = (
            db.query(Task)
            .filter(Task.document_id.in_(doc_ids))
            .order_by(Task.created_at.asc())  # asc so latest overwrites earlier
            .all()
        )
        for t in tasks:
            task_summary_map[t.document_id][t.task_type] = t.status

    return [
        DocumentListItem(
            id=d.id,
            title=d.title,
            original_filename=d.original_filename,
            format=d.format,
            total_pages=d.total_pages,
            processing_status=d.processing_status,
            source_language=d.source_language,
            created_at=d.created_at.isoformat() if d.created_at else None,
            task_summary=task_summary_map.get(d.id) or None,
        )
        for d in docs
    ]


# ── Get document detail ─────────────────────────────────────────────

@router.get("/{document_id}", response_model=DocumentDetailResponse)
async def get_document(
    document_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    Get document metadata.

    Access rules: users can only access their own documents.
    ADMIN can access any document.
    """
    repo = DocumentRepository(db)
    doc = repo.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if user.role != "ADMIN" and doc.user_id != user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    return DocumentDetailResponse(
        id=doc.id,
        title=doc.title,
        original_filename=doc.original_filename,
        source_language=doc.source_language,
        format=doc.format,
        file_type=doc.file_type,
        total_pages=doc.total_pages,
        processing_status=doc.processing_status,
        user_id=doc.user_id,
        created_at=doc.created_at.isoformat() if doc.created_at else None,
        updated_at=doc.updated_at.isoformat() if doc.updated_at else None,
    )


# ── Get document text (OCR / normalized) ────────────────────────────

@router.get("/{document_id}/text", response_model=DocumentTextResponse)
async def get_document_text(
    document_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Get OCR and/or normalized text for a document (own docs only; admin can access all)."""
    repo = DocumentRepository(db)
    doc = repo.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if user.role != "ADMIN" and doc.user_id != user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    dt = repo.get_digitized_text(document_id)
    return DocumentTextResponse(
        document_id=document_id,
        ocr_content=dt.ocr_content if dt else None,
        normalized_content=dt.normalized_content if dt else None,
    )


# ── Download OCR / normalized text as file ──────────────────────────

@router.get("/{document_id}/text/download")
async def download_document_text(
    document_id: str,
    type: str = Query("ocr", pattern="^(ocr|normalized)$"),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    Download extracted text as a .docx file.

    - ?type=ocr        → raw OCR content (default)
    - ?type=normalized → cleaned/normalized content
    """
    import io
    from docx import Document as DocxDocument

    repo = DocumentRepository(db)
    doc = repo.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if user.role != "ADMIN" and doc.user_id != user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    dt = repo.get_digitized_text(document_id)
    if not dt:
        raise HTTPException(status_code=404, detail="No extracted text found. Run /extract first.")

    content = dt.ocr_content if type == "ocr" else dt.normalized_content
    if not content:
        raise HTTPException(status_code=404, detail=f"No {type} content available.")

    docx = DocxDocument()
    docx.add_heading(doc.title, level=1)
    for line in content.splitlines():
        docx.add_paragraph(line)

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)

    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in doc.title)[:60]
    filename = f"{type}_{safe_title}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Upload corrected OCR / text ──────────────────────────────────────

@router.post("/{document_id}/text/upload", response_model=DocumentTextResponse)
async def upload_document_text(
    document_id: str,
    file: UploadFile = File(..., description="Corrected document text as .txt or .docx"),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    Override the extracted/OCR text by uploading a corrected .txt or .docx file.
    Writes to normalized_content (the cleaned text used by all downstream pipeline steps).
    """
    repo = DocumentRepository(db)
    doc = repo.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if user.role != "ADMIN" and doc.user_id != user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    if not repo.get_digitized_text(document_id):
        raise HTTPException(
            status_code=409,
            detail="No extracted text exists yet. Run /extract first, then upload your correction.",
        )

    text = await extract_text_from_upload(file)
    dt = repo.update_digitized_text(document_id, text)
    return DocumentTextResponse(
        document_id=document_id,
        ocr_content=dt.ocr_content,
        normalized_content=dt.normalized_content,
    )


# ── Get pages with markdown ─────────────────────────────────────────

@router.get("/{document_id}/pages", response_model=list[PageListItem])
async def get_document_pages(
    document_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Get all pages for a document (own docs only; admin can access all)."""
    repo = DocumentRepository(db)
    doc = repo.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if user.role != "ADMIN" and doc.user_id != user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    pages = repo.get_pages(document_id)
    return [
        PageListItem(
            id=p.id,
            page_number=p.page_number,
            markdown_content=p.markdown_content,
            image_width=p.image_width,
            image_height=p.image_height,
        )
        for p in pages
    ]


# ── Get layout elements ─────────────────────────────────────────────

@router.get("/{document_id}/elements", response_model=list[ElementListItem])
async def get_document_elements(
    document_id: str,
    label: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Get layout elements (bounding boxes) for a document (own docs only; admin can access all)."""
    repo = DocumentRepository(db)
    doc = repo.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if user.role != "ADMIN" and doc.user_id != user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    elements = repo.get_elements(document_id, label=label)
    result = []
    for elem in elements:
        # Page is already joined; fetch page_number via page relationship or separate query
        from data.db_models import Page as PageModel
        page = db.query(PageModel).filter(PageModel.id == elem.page_id).first()
        result.append(ElementListItem(
            id=elem.id,
            label=elem.label,
            text_content=elem.text_content,
            bbox={
                "x1": elem.bbox_x1,
                "y1": elem.bbox_y1,
                "x2": elem.bbox_x2,
                "y2": elem.bbox_y2,
            },
            bbox_normalized={
                "x1": elem.bbox_norm_x1,
                "y1": elem.bbox_norm_y1,
                "x2": elem.bbox_norm_x2,
                "y2": elem.bbox_norm_y2,
            } if elem.bbox_norm_x1 is not None else None,
            page_number=page.page_number if page else None,
            page_id=elem.page_id,
            sequence_order=elem.sequence_order,
            has_crop_image=bool(elem.crop_image_base64),
        ))
    return result


# ── Delete document ─────────────────────────────────────────────────

@router.delete("/{document_id}", status_code=204)
async def delete_document(
    document_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    Permanently delete a document and all its related data (pages, OCR text,
    translations, summaries, keywords, research directions, tasks).

    Access: ADMIN can delete any document; regular users can only delete their own.
    The uploaded file on disk is also removed.
    """
    repo = DocumentRepository(db)
    doc = repo.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if user.role != "ADMIN" and doc.user_id != user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    # Remove file from disk (best-effort — don't fail the request if missing)
    if doc.file_path and os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except OSError:
            pass

    db.delete(doc)
    db.commit()
