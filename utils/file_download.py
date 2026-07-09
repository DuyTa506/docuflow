"""Shared helper for building .docx / .pdf download responses."""

import io
import mimetypes
import os
import tempfile
import unicodedata
from pathlib import Path
from urllib.parse import quote

from docx import Document as DocxDocument
from fastapi import HTTPException, Response
from fastapi.responses import StreamingResponse

from config.settings import settings
from utils.markdown_docx import build_docx_bytes_from_markdown, render_layout_elements_to_docx
from utils.markdown_pandoc import markdown_to_docx_bytes
from utils.ocr_markdown import is_structured_markdown, normalize_ocr_markdown, sanitize_xml_text

_NATIVE_WORD_FORMATS = frozenset({"docx", "doc"})


def is_native_word_document(doc_format: str | None) -> bool:
    return (doc_format or "").lower() in _NATIVE_WORD_FORMATS


def build_stored_file_response(
    storage_key: str,
    *,
    download_name: str | None = None,
    content_type: str | None = None,
) -> StreamingResponse:
    """Stream a file from MinIO object storage."""
    from services.object_storage import get_object_storage

    storage = get_object_storage()
    if not storage_key or not storage.exists(storage_key):
        raise HTTPException(status_code=404, detail="File not found in storage.")

    filename = download_name or os.path.basename(storage_key)
    media_type = content_type or mimetypes.guess_type(filename)[0] or "application/octet-stream"
    encoded = quote(filename, safe="")
    disposition = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{encoded}'
    return StreamingResponse(
        storage.iter_stream(storage_key),
        media_type=media_type,
        headers={"Content-Disposition": disposition},
    )


def build_docx_response(
    filename: str,
    content: str,
    *,
    title: str | None = None,
    headings: list[str] | None = None,
    structured: bool = True,
    engine: str | None = None,
) -> Response:
    """Build a .docx Response from text content."""
    export_engine = engine or settings.docx_export_engine

    if structured and is_structured_markdown(content) and export_engine != "python":
        body = markdown_to_docx_bytes(
            normalize_ocr_markdown(content),
            title=title,
            headings=headings,
        )
    elif structured and is_structured_markdown(content):
        body = build_docx_bytes_from_markdown(
            normalize_ocr_markdown(content),
            title=title,
            headings=headings,
        )
    else:
        docx = DocxDocument()
        if title:
            docx.add_heading(title, level=1)
        for h in headings or []:
            docx.add_heading(h, level=2)
        for line in content.splitlines():
            docx.add_paragraph(line)
        buf = io.BytesIO()
        docx.save(buf)
        body = buf.getvalue()

    return _docx_bytes_response(filename, body)


def build_docx_bytes_from_content(
    content: str,
    *,
    title: str | None = None,
    headings: list[str] | None = None,
    structured: bool = True,
) -> bytes:
    """Return raw .docx bytes (for PDF conversion pipeline)."""
    content = sanitize_xml_text(content)
    title = sanitize_xml_text(title) if title else None
    headings = [sanitize_xml_text(h) for h in (headings or [])]
    if structured and is_structured_markdown(content):
        return markdown_to_docx_bytes(
            normalize_ocr_markdown(content),
            title=title,
            headings=headings,
        )
    docx = DocxDocument()
    if title:
        docx.add_heading(title, level=1)
    for h in headings or []:
        docx.add_heading(h, level=2)
    for line in content.splitlines():
        docx.add_paragraph(line)
    buf = io.BytesIO()
    docx.save(buf)
    return buf.getvalue()


def build_docx_bytes_from_elements(
    elements, *, title: str | None = None, document_id: str | None = None, embed_images: bool = True
) -> bytes:
    from utils.translation_elements import elements_to_views

    docx = DocxDocument()
    if title:
        docx.add_heading(title, level=1)
    render_layout_elements_to_docx(
        docx,
        elements_to_views(elements, document_id=document_id),
        embed_images=embed_images,
    )
    buf = io.BytesIO()
    docx.save(buf)
    return buf.getvalue()


def build_pdf_bytes_from_elements(
    elements,
    pages,
    *,
    document_id: str | None = None,
    page_background: bool | None = None,
    merge_blocks: bool = False,
    text_overlay: str = "skip",
    page_backgrounds: dict | None = None,
) -> bytes:
    """Build layout-faithful PDF (one source page → one PDF page)."""
    from utils.layout_pdf import build_layout_pdf_bytes
    from utils.translation_blocks import merge_elements_for_layout_export
    from utils.translation_elements import layout_element_to_dict

    export_elements = elements
    if merge_blocks:
        payloads: list = []
        for elem in elements:
            if isinstance(elem, dict):
                payloads.append(elem)
            else:
                page_num = 1
                if getattr(elem, "page", None) is not None:
                    page_num = getattr(elem.page, "page_number", 1) or 1
                payloads.append(layout_element_to_dict(elem, page_num))
        export_elements = merge_elements_for_layout_export(payloads)

    return build_layout_pdf_bytes(
        export_elements,
        pages,
        document_id=document_id,
        page_background=page_background,
        text_overlay=text_overlay,  # type: ignore[arg-type]
        page_backgrounds=page_backgrounds,
    )


def docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF via LibreOffice headless."""
    from utils.soffice import run_soffice

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        src = tmp / "input.docx"
        src.write_bytes(docx_bytes)
        result = run_soffice(
            ["--headless", "--convert-to", "pdf", "--outdir", str(tmp), str(src)],
            capture_output=True,
            text=True,
            timeout=180,
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
        pdf_path = tmp / "input.pdf"
        if not pdf_path.exists():
            raise RuntimeError("LibreOffice produced no PDF output")
        return pdf_path.read_bytes()


def _docx_bytes_response(filename: str, body: bytes) -> Response:
    encoded = quote(filename, safe="")
    disposition = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{encoded}'
    return Response(
        content=body,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": disposition},
    )


def safe_filename(text: str, max_len: int = 60) -> str:
    """Sanitize a string for use in a filename, keeping only ASCII alphanumeric chars."""
    text = unicodedata.normalize("NFKD", text)
    text = "".join(c for c in text if not unicodedata.combining(c))
    return "".join(c if c.isascii() and (c.isalnum() or c in " -_") else "_" for c in text)[:max_len]
