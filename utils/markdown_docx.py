"""
Convert OCR markdown (with embedded HTML tables) into a structured python-docx document.
"""

from __future__ import annotations

import re
from typing import Iterable

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph

from core.constants import OCR_EQUATION_LABELS
from utils.ocr_markdown import normalize_ocr_markdown, split_pages
from utils.table_grid import (
    HtmlTableParser as _HtmlTableParser,
    build_table_grid as _build_table_grid,
    compact_empty_columns as _compact_empty_columns,
    parse_markdown_table_rows as _parse_markdown_table_rows,
)

_IMAGE_LABELS = frozenset({"image", "figure", "chart", "graph", "picture"})
_IMG_PLACEHOLDER_RE = re.compile(r"^\(img_content\)|^\[?image[_\s]?\d*\]?$|^\[figure", re.IGNORECASE)
_USABLE_PAGE_WIDTH_IN = 6.3

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_UL_RE = re.compile(r"^(\s*)[-*+]\s+(.+)$")
_OL_RE = re.compile(r"^(\s*)\d+\.\s+(.+)$")
_TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")
_HTML_TABLE_RE = re.compile(r"(?is)<table\b.*?</table>")
_CODE_FENCE_RE = re.compile(r"^```")

# Vietnamese administrative/academic document formatting convention
# (matches Circular 01/2011/TT-BNV-style expectations): Times New Roman,
# body text at 14pt, headings bold and stepped down in size per level.
_BODY_FONT = "Times New Roman"
_BODY_SIZE_PT = 14
_HEADING_SIZES_PT = {1: 16, 2: 15, 3: 14, 4: 14}
_BODY_FIRST_LINE_INDENT_CM = 1.0


def _nsdecl() -> str:
    return 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _apply_document_defaults(doc: DocxDocument) -> None:
    """Set Times New Roman 14pt as the document's base style, and size/bold
    the built-in Heading 1-4 styles per Vietnamese document convention,
    instead of leaving whatever font python-docx's default template ships
    with (Calibri) or reconstructing the source PDF's original visual
    layout -- exported DOCX should read as a normal, standard document,
    not a geometric clone of the source."""
    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = Pt(_BODY_SIZE_PT)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {_nsdecl()} w:ascii="{_BODY_FONT}" w:hAnsi="{_BODY_FONT}" w:eastAsia="{_BODY_FONT}" w:cs="{_BODY_FONT}"/>')
        rpr.append(rfonts)
    else:
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(attr), _BODY_FONT)
    normal.paragraph_format.first_line_indent = Cm(_BODY_FIRST_LINE_INDENT_CM)
    normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for level, size_pt in _HEADING_SIZES_PT.items():
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.name = _BODY_FONT
        style.font.size = Pt(size_pt)
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER if level == 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
        )


def render_markdown_to_docx(
    doc: DocxDocument,
    markdown: str,
    *,
    page_breaks: bool = True,
) -> None:
    """Append structured content from OCR markdown into an existing document."""
    _apply_document_defaults(doc)
    markdown = normalize_ocr_markdown(markdown)
    if not markdown:
        return

    pages = split_pages(markdown) if page_breaks and "\n---\n" in f"\n{markdown}\n" else [markdown]
    for idx, page_text in enumerate(pages):
        if idx > 0 and page_breaks:
            doc.add_page_break()
        _render_block(doc, page_text)


def render_layout_elements_to_docx(
    doc: DocxDocument,
    elements: Iterable,
    *,
    page_break_between_pages: bool = True,
    embed_images: bool = True,
) -> None:
    """Render document from ordered layout elements (spatial reading order)."""
    from utils.ocr_markdown import element_export_text, element_heading_level

    _apply_document_defaults(doc)

    views = list(elements)
    page_metrics = _build_page_metrics(views)

    current_page: int | None = None
    prev_bottom: float | None = None
    for elem in views:
        page_num = getattr(elem, "page_number", None)
        page_rel = getattr(elem, "page", None)
        if page_num is None and page_rel is not None:
            page_num = getattr(page_rel, "page_number", None)

        if (
            page_break_between_pages
            and page_num is not None
            and current_page is not None
            and page_num != current_page
        ):
            doc.add_page_break()
        if page_num != current_page:
            prev_bottom = None  # reset vertical tracking on a new page
        if page_num is not None:
            current_page = page_num

        metrics = page_metrics.get(page_num)
        before = len(doc.paragraphs)
        _render_element(doc, elem, metrics, embed_images=embed_images)

        space = _vertical_space_before(elem, prev_bottom, metrics)
        if space is not None and len(doc.paragraphs) > before:
            doc.paragraphs[before].paragraph_format.space_before = space

        y2 = getattr(elem, "bbox_y2", None)
        if y2 is not None:
            prev_bottom = float(y2)


def _render_element(doc: DocxDocument, elem, metrics: "_PageMetrics | None", *, embed_images: bool = True) -> None:
    """Render a single layout element (image / heading / table / equation / text)."""
    from utils.ocr_markdown import element_export_text, element_heading_level

    label = getattr(elem, "label", "text") or "text"

    if embed_images and label.lower() in _IMAGE_LABELS:
        if _add_image_paragraph(doc, elem, metrics):
            return
        # fall through to render any caption / placeholder text

    text = element_export_text(label, getattr(elem, "text_content", None))
    if not text:
        return

    level = element_heading_level(label)
    if level is not None:
        # Alignment/size/bold now come from the "Heading N" style itself
        # (set once in _apply_document_defaults) rather than the source
        # element's bbox position -- a standard document shouldn't visually
        # clone the original PDF/scan's geometry.
        doc.add_heading(_strip_inline_markdown(text), level=level)
        return

    if label.lower() == "table" or "<table" in text.lower():
        _render_tables_and_text(doc, text)
        return

    if label.lower() in OCR_EQUATION_LABELS:
        _add_equation_paragraph(doc, text, elem, metrics)
        return

    p = doc.add_paragraph()
    _add_multiline_runs(p, text)
    # Justify + first-line indent come from the "Normal" style default
    # (Vietnamese document convention) -- no per-paragraph override needed.


def build_docx_bytes_from_markdown(
    markdown: str,
    *,
    title: str | None = None,
    headings: list[str] | None = None,
) -> bytes:
    """Build a complete .docx byte stream from markdown content."""
    import io

    doc = DocxDocument()
    if title:
        doc.add_heading(title, level=1)
    for h in headings or []:
        doc.add_heading(h, level=2)
    render_markdown_to_docx(doc, markdown)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_block(doc: DocxDocument, text: str) -> None:
    remaining = text.strip()
    if not remaining:
        return

    while remaining:
        html_match = _HTML_TABLE_RE.search(remaining)
        if html_match and html_match.start() == 0:
            table_html = html_match.group(0)
            _add_html_table(doc, table_html)
            remaining = remaining[html_match.end() :].lstrip()
            continue

        if html_match and html_match.start() > 0:
            prefix = remaining[: html_match.start()].rstrip()
            if prefix:
                _render_text_lines(doc, prefix)
            remaining = remaining[html_match.start() :].lstrip()
            continue

        _render_text_lines(doc, remaining)
        break


def _render_tables_and_text(doc: DocxDocument, text: str) -> None:
    _render_block(doc, text)


def _render_text_lines(doc: DocxDocument, text: str) -> None:
    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if _CODE_FENCE_RE.match(stripped):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
                run.font.size = Pt(10)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if _is_markdown_table_start(lines, i):
            table_lines, i = _collect_markdown_table(lines, i)
            _add_markdown_table(doc, table_lines)
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = min(len(heading.group(1)), 9)
            doc.add_heading(_strip_inline_markdown(heading.group(2)), level=level)
            i += 1
            continue

        ul = _UL_RE.match(line)
        if ul:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, ul.group(2).strip())
            i += 1
            continue

        ol = _OL_RE.match(line)
        if ol:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, ol.group(2).strip())
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or _HEADING_RE.match(nxt)
                or _UL_RE.match(lines[i])
                or _OL_RE.match(lines[i])
                or _is_markdown_table_start(lines, i)
                or _HTML_TABLE_RE.match(nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1

        if _should_center_line_block(para_lines):
            for line in para_lines:
                p = doc.add_paragraph()
                _add_inline_runs(p, line)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(para_lines))


def _is_markdown_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    row = lines[index].strip()
    sep = lines[index + 1].strip()
    return bool(_TABLE_ROW_RE.match(row) and _TABLE_SEP_RE.match(sep))


def _collect_markdown_table(lines: list[str], start: int) -> tuple[list[str], int]:
    collected = []
    i = start
    while i < len(lines):
        row = lines[i].strip()
        if not row or not _TABLE_ROW_RE.match(row):
            break
        collected.append(row)
        i += 1
    return collected, i


def _add_markdown_table(doc: DocxDocument, table_lines: list[str]) -> None:
    str_rows = _parse_markdown_table_rows(table_lines)
    _add_rows_to_docx_table(doc, str_rows)


def _add_html_table(doc: DocxDocument, html: str) -> None:
    parser = _HtmlTableParser()
    parser.feed(html)
    parser.close()
    if parser.rows:
        _render_grid_table(doc, parser.rows, table_style=parser.table_style)


def _add_rows_to_docx_table(doc: DocxDocument, rows: list[list[str]]) -> None:
    """Compatibility wrapper: render plain string rows (no spans) via the grid."""
    if not rows:
        return
    cell_rows: list[list[dict]] = []
    for r_idx, row in enumerate(rows):
        cell_rows.append(
            [
                {"text": text, "colspan": 1, "rowspan": 1, "header": r_idx == 0}
                for text in row
            ]
        )
    _render_grid_table(doc, cell_rows)


def _render_grid_table(
    doc: DocxDocument,
    rows: list[list[dict]],
    *,
    table_style: str = "",
) -> None:
    n_rows, n_cols, placements = _build_table_grid(rows)
    n_cols, placements = _compact_empty_columns(n_cols, placements)
    if n_rows == 0 or n_cols == 0 or not placements:
        return

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    if "margin: auto" in table_style or "margin:auto" in table_style:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for (r0, c0, r1, c1, text, header) in placements:
        try:
            origin = table.cell(r0, c0)
            if (r0, c0) != (r1, c1):
                origin = origin.merge(table.cell(r1, c1))
        except (IndexError, ValueError):
            continue
        _set_grid_cell(origin, text, header or r0 == 0)


def _set_grid_cell(cell, text: str, header: bool) -> None:
    """Write rich text into a (possibly merged) table cell.

    Converts literal ``\\n`` and real newlines into line breaks; renders inline
    markdown / math via _add_inline_runs; bolds header cells.
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]
    normalized = (
        _strip_html_inline(text or "")
        .replace("\\r\\n", "\n")
        .replace("\\n", "\n")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )
    lines = normalized.split("\n")
    while lines and not lines[-1].strip():
        lines.pop()
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run().add_break()
        _add_inline_runs(paragraph, line)
    if header:
        for run in paragraph.runs:
            run.bold = True


_MATH_SPLIT_RE = re.compile(r"(\$\$.+?\$\$|(?<!\$)\$(?!\$).+?(?<!\$)\$(?!\$))", re.DOTALL)


def _add_inline_runs(paragraph: Paragraph, text: str) -> None:
    """Render inline content, converting ``$...$`` / ``$$...$$`` math segments."""
    from utils.math_omml import looks_like_math

    text = _strip_html_inline(text)
    if not text:
        return

    for part in _MATH_SPLIT_RE.split(text):
        if not part:
            continue
        if part.startswith("$$") and part.endswith("$$") and len(part) > 4:
            inner = part[2:-2]
            if looks_like_math(inner):
                _add_math_run(paragraph, inner, display=True)
                continue
        elif part.startswith("$") and part.endswith("$") and len(part) > 2:
            inner = part[1:-1]
            if looks_like_math(inner):
                _add_math_run(paragraph, inner, display=False)
                continue
        _add_markdown_runs(paragraph, part)


def _add_markdown_runs(paragraph: Paragraph, text: str) -> None:
    if not text:
        return
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_|`[^`]+`)"
    )
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("_") and part.endswith("_"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        else:
            paragraph.add_run(part)


def _add_math_run(paragraph: Paragraph, latex: str, *, display: bool) -> None:
    """Insert a math fragment: native OMML for real LaTeX, else superscript text."""
    from utils.math_omml import (
        has_latex_command,
        latex_to_omml_fragment,
        omml_fragment_for_docx,
        sanitize_latex_fragment,
    )

    body = sanitize_latex_fragment(latex)
    if not body:
        return
    if has_latex_command(body):
        omml = latex_to_omml_fragment(body, display=display)
        if omml:
            try:
                paragraph._element.append(parse_xml(omml_fragment_for_docx(omml)))
                return
            except Exception:
                pass
    _add_latex_text_fallback(paragraph, body)


def _add_latex_text_fallback(paragraph: Paragraph, latex: str) -> None:
    """Render simple LaTeX as Word runs: ``^{..}`` -> superscript, ``_{..}`` -> subscript."""
    from utils.math_omml import replace_latex_commands

    s = replace_latex_commands(latex)
    buf = ""
    i = 0

    def _flush():
        nonlocal buf
        if buf:
            paragraph.add_run(buf)
            buf = ""

    while i < len(s):
        ch = s[i]
        if ch in "^_" and i + 1 < len(s):
            _flush()
            i += 1
            if s[i] == "{":
                j = s.find("}", i)
                if j == -1:
                    content, i = s[i + 1:], len(s)
                else:
                    content, i = s[i + 1:j], j + 1
            else:
                content, i = s[i], i + 1
            if content:
                run = paragraph.add_run(content)
                if ch == "^":
                    run.font.superscript = True
                else:
                    run.font.subscript = True
        elif ch in "{}":
            i += 1
        else:
            buf += ch
            i += 1
    _flush()


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"^#{1,6}\s+", "", text.strip())
    return _strip_html_inline(text)


def _strip_html_inline(text: str) -> str:
    text = re.sub(r"(?i)</?center>", "", text)
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    return text.strip()


def _should_center_line_block(lines: list[str]) -> bool:
    """Heuristic: title-page style blocks — several short lines, no sentence end."""
    if len(lines) < 2:
        return False
    if any(len(line) > 90 for line in lines):
        return False
    if any(line.endswith((".", "?", "!")) and len(line) > 40 for line in lines):
        return False
    return True


def _add_multiline_runs(paragraph: Paragraph, text: str) -> None:
    """Preserve intentional line breaks inside a spatial text block."""
    text = _strip_html_inline(text)
    if not text:
        return
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return
    if len(lines) == 1:
        _add_inline_runs(paragraph, lines[0])
        return
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run().add_break()
        _add_inline_runs(paragraph, line)


class _PageMetrics:
    __slots__ = ("width", "min_x", "max_x", "height", "min_y", "max_y")

    def __init__(
        self,
        width: float,
        min_x: float,
        max_x: float,
        height: float = 0.0,
        min_y: float = 0.0,
        max_y: float = 0.0,
    ):
        self.width = width
        self.min_x = min_x
        self.max_x = max_x
        self.height = height
        self.min_y = min_y
        self.max_y = max_y


def _build_page_metrics(elements: Iterable) -> dict[int, _PageMetrics]:
    buckets: dict[int, list[tuple[float, float, float, float]]] = {}
    for elem in elements:
        page_num = getattr(elem, "page_number", None)
        x1 = getattr(elem, "bbox_x1", None)
        x2 = getattr(elem, "bbox_x2", None)
        if page_num is None or x1 is None or x2 is None:
            continue
        y1 = getattr(elem, "bbox_y1", None)
        y2 = getattr(elem, "bbox_y2", None)
        buckets.setdefault(page_num, []).append(
            (float(x1), float(x2), float(y1) if y1 is not None else 0.0,
             float(y2) if y2 is not None else 0.0)
        )

    metrics: dict[int, _PageMetrics] = {}
    for page_num, boxes in buckets.items():
        min_x = min(b[0] for b in boxes)
        max_x = max(b[1] for b in boxes)
        min_y = min(b[2] for b in boxes)
        max_y = max(b[3] for b in boxes)
        width = max(max_x - min_x, 1.0)
        height = max(max_y - min_y, 1.0)
        metrics[page_num] = _PageMetrics(
            width=width, min_x=min_x, max_x=max_x,
            height=height, min_y=min_y, max_y=max_y,
        )
    return metrics


def _vertical_space_before(elem, prev_bottom: float | None, metrics: _PageMetrics | None):
    """Add paragraph space_before proportional to the vertical gap from the
    previous element (reconstructs intentional whitespace). Conservative: only
    fires on gaps above ~5% of page height, capped at 28pt."""
    if metrics is None or prev_bottom is None or metrics.height <= 0:
        return None
    y1 = getattr(elem, "bbox_y1", None)
    if y1 is None:
        return None
    gap = float(y1) - float(prev_bottom)
    if gap <= 0:
        return None
    frac = gap / metrics.height
    if frac < 0.05:
        return None
    return Pt(min(frac * 792.0, 28.0))


def _load_element_image_bytes(elem) -> bytes | None:
    """Resolve image bytes for a figure/image element.

    Priority: MinIO crop -> inline base64 crop -> crop from full page render.
    """
    import base64 as _b64

    key = getattr(elem, "crop_image_key", None)
    if key:
        try:
            from services.object_storage import get_object_storage

            return get_object_storage().get_bytes(key)
        except Exception:
            pass

    b64 = getattr(elem, "crop_image_base64", None)
    if b64:
        try:
            return _b64.b64decode(b64)
        except Exception:
            pass

    return _crop_from_page_image(elem)


def _crop_from_page_image(elem) -> bytes | None:
    """Crop the element's bbox region out of the stored full-page image."""
    page_key = getattr(elem, "page_image_key", None)
    x1 = getattr(elem, "bbox_x1", None)
    y1 = getattr(elem, "bbox_y1", None)
    x2 = getattr(elem, "bbox_x2", None)
    y2 = getattr(elem, "bbox_y2", None)
    if not page_key or x1 is None or y1 is None or x2 is None or y2 is None:
        return None
    try:
        from io import BytesIO

        from PIL import Image

        from services.object_storage import get_object_storage

        data = get_object_storage().get_bytes(page_key)
        img = Image.open(BytesIO(data))
        box = (int(x1), int(y1), int(x2), int(y2))
        if box[2] <= box[0] or box[3] <= box[1]:
            return None
        crop = img.crop(box)
        buf = BytesIO()
        crop.convert("RGB").save(buf, format="JPEG", quality=90)
        return buf.getvalue()
    except Exception:
        return None


def _image_width_for_elem(elem, metrics: _PageMetrics | None):
    """Scale the picture width to the element's bbox fraction of the page width."""
    x1 = getattr(elem, "bbox_x1", None)
    x2 = getattr(elem, "bbox_x2", None)
    if metrics is None or x1 is None or x2 is None or metrics.width <= 0:
        return None
    frac = (float(x2) - float(x1)) / metrics.width
    frac = max(0.1, min(frac, 1.0))
    return Inches(min(frac * _USABLE_PAGE_WIDTH_IN, _USABLE_PAGE_WIDTH_IN))


def _add_image_paragraph(doc: DocxDocument, elem, metrics: _PageMetrics | None) -> bool:
    """Embed a figure/image element as a real picture. Returns True on success."""
    from io import BytesIO

    img_bytes = _load_element_image_bytes(elem)
    if not img_bytes:
        return False
    try:
        p = doc.add_paragraph()
        run = p.add_run()
        width = _image_width_for_elem(elem, metrics)
        if width is not None:
            run.add_picture(BytesIO(img_bytes), width=width)
        else:
            run.add_picture(BytesIO(img_bytes))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception:
        return False

    caption = _strip_html_inline(getattr(elem, "text_content", "") or "").strip()
    if caption and not _IMG_PLACEHOLDER_RE.match(caption):
        cap = doc.add_paragraph()
        _add_inline_runs(cap, caption)
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(9)
        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return True


def _add_equation_paragraph(doc: DocxDocument, text: str, elem, metrics: _PageMetrics | None) -> None:
    """Render equation as OMML when pandoc is available, else superscript-aware text."""
    from utils.math_omml import (
        latex_to_omml_fragment,
        omml_fragment_for_docx,
        sanitize_latex_fragment,
        wrap_as_equation_markdown,
    )

    latex = wrap_as_equation_markdown(text).strip()
    inner = sanitize_latex_fragment(latex.strip("$").strip())
    omml_bytes = latex_to_omml_fragment(inner, display=True) if inner else None
    p = doc.add_paragraph()
    if omml_bytes:
        try:
            p._element.append(parse_xml(omml_fragment_for_docx(omml_bytes)))
        except Exception:
            _add_latex_text_fallback(p, inner or text)
    else:
        _add_latex_text_fallback(p, inner or text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # standard convention for display equations


