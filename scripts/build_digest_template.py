#!/usr/bin/env python3
"""Generate docuflow_digest_template.docx with docxtpl Jinja2 placeholders."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "template" / "docuflow_digest_template.docx"


def _sep(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("*****")
    r.bold = True


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h1.add_run("HỌC VIỆN KỸ THUẬT QUÂN SỰ")
    r.bold = True
    r.font.size = Pt(13)

    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = h2.add_run("PHÒNG THÔNG TIN KHOA HỌC QUÂN SỰ")
    r2.bold = True

    _sep(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run("TỔNG THUẬT TÀI LIỆU")
    rt.bold = True
    rt.font.size = Pt(14)

    doc.add_paragraph()

    # §1
    p = doc.add_paragraph()
    p.add_run("1. THÔNG TIN CHUNG VỀ TÀI LIỆU").bold = True

    # Labels are quoted verbatim from `test_docs/Mau Tong thuat Book .docx` —
    # the bilingual suffix appears on some and not others, and that is the mẫu's
    # choice, not an inconsistency to tidy up.
    for label, key in [
        ("Nhan đề (Title)", "title_display"),
        ("Tác giả (Authors)", "authors"),
        ("Nhà xuất bản (Publisher)", "publisher"),
        ("Năm xuất bản (Year)", "year"),
        ("ISBN", "isbn"),
        ("DOI", "doi"),
        ("Số trang", "pages"),
    ]:
        fp = doc.add_paragraph()
        fp.add_run(f"- {label}: ").bold = True
        fp.add_run(f"{{{{ bib.{key} }}}}")

    doc.add_paragraph()

    # §2
    p2 = doc.add_paragraph()
    p2.add_run("2. TỔNG THUẬT VỀ TÀI LIỆU").bold = True

    # `{%p ... %}` is docxtpl's paragraph-aware form: the paragraph holding the
    # tag is consumed. Plain `{% ... %}` leaves it behind as a blank line, one
    # per loop iteration.
    s21 = doc.add_paragraph()
    s21.add_run("2.1. Tóm tắt").bold = True
    # `{{p x }}` inserts a subdoc at paragraph level. The renderer builds those
    # subdocs with the same inline renderer the text-download path uses, so
    # `$…$` becomes a real Word equation and `**bold**` a bold run instead of
    # printing the markup.
    doc.add_paragraph("{%p for para in abstract_paragraphs %}")
    doc.add_paragraph("{{p para }}")
    doc.add_paragraph("{%p endfor %}")

    doc.add_paragraph()
    s22 = doc.add_paragraph()
    s22.add_run("2.2. Nội dung chính của tài liệu").bold = True
    doc.add_paragraph("{%p for ch in chapters %}")
    # `body` is ONE paragraph carrying label + summary together, as the mẫu
    # requires (`Chương 1. Giới thiệu (Введение). Nội dung…`). Composed in
    # DigestRenderer so the three label forms (Chương N. / cụm BBKH / BBKH N -)
    # stay one tested function instead of Jinja branches inside Word XML.
    doc.add_paragraph("{{p ch.body }}")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p if not chapters %}")
    doc.add_paragraph("[Chưa có — chạy main_content trước]")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph()
    s23 = doc.add_paragraph()
    s23.add_run("2.3. Từ khóa").bold = True
    doc.add_paragraph("{%p for kw in keywords %}")
    doc.add_paragraph("- {{ kw.display }}")
    doc.add_paragraph("{%p endfor %}")

    doc.add_paragraph()

    # §3
    p3 = doc.add_paragraph()
    p3.add_run("3. PHẠM VI SỬ DỤNG").bold = True

    for label, key in [
        ("CTĐT đại học", "undergraduate_text"),
        ("CTĐT thạc sĩ", "master_text"),
        ("CTĐT tiến sĩ", "phd_text"),
        ("Nhóm nghiên cứu mạnh", "strong_research_groups_text"),
    ]:
        fp = doc.add_paragraph()
        fp.add_run(f"- {label}: ").bold = True
        fp.add_run(f"{{{{ usage.{key} }}}}")

    rd = doc.add_paragraph()
    rd.add_run("- Hướng nghiên cứu: ").bold = True
    rd.add_run("{{ research_directions_text }}")

    doc.add_paragraph()
    _sep(doc)

    admin = doc.add_paragraph()
    admin.add_run("* Thông tin quản trị CSDL:").bold = True
    for label, key in [
        ("Người tổng thuật", "reviewer"),
        ("Người hiệu đính, phê duyệt", "reviewer_approved"),
        ("Ngày nhập CSDL Học liệu số", "entry_date"),
    ]:
        fp = doc.add_paragraph()
        fp.add_run(f"- {label}: ").bold = True
        fp.add_run(f"{{{{ {key} }}}}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
