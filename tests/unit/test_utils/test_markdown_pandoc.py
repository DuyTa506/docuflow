"""Tests for pandoc-based markdown export."""

import shutil

import pytest
from docx import Document as DocxDocument

from utils.markdown_pandoc import is_pandoc_available, markdown_to_docx_bytes
from utils.ocr_markdown import element_export_text, is_structured_markdown


@pytest.mark.skipif(not is_pandoc_available(), reason="pandoc not installed")
class TestMarkdownPandoc:
    def test_latex_produces_omml(self):
        md = "Energy: $E = mc^2$ and display:\n\n$$\\int_0^1 x\\,dx$$\n"
        assert is_structured_markdown(md) is True
        body = markdown_to_docx_bytes(md, title="Math test")
        doc = DocxDocument(__import__("io").BytesIO(body))
        xml = doc.element.body.xml
        assert "oMath" in xml or "oMathPara" in xml


class TestMarkdownPandocFallback:
    def test_fallback_without_pandoc(self, monkeypatch):
        monkeypatch.setattr(shutil, "which", lambda _: None)
        assert is_pandoc_available() is False
        body = markdown_to_docx_bytes("# Hello\n\nPlain text.", title="T")
        assert len(body) > 100
        doc = DocxDocument(__import__("io").BytesIO(body))
        assert any(p.text for p in doc.paragraphs)

    @pytest.mark.skipif(is_pandoc_available(), reason="pandoc installed; fallback covered above")
    def test_works_when_pandoc_not_on_path(self):
        body = markdown_to_docx_bytes("## Section\n\nBody text.", title="No pandoc")
        assert len(body) > 100


class TestOcrMarkdownEquations:
    def test_equation_label_wrapped(self):
        out = element_export_text("equation", "E = mc^2")
        assert out.startswith("$$") and out.endswith("$$")

    def test_isolate_formula_label_wrapped(self):
        out = element_export_text("isolate_formula", r"\int_0^1 x\,dx")
        assert out.startswith("$$") and out.endswith("$$")

    def test_bracket_latex_normalized(self):
        out = element_export_text("equation", r"\[ E=mc^{2} \]")
        assert out == "$$E=mc^{2}$$"

    def test_detects_latex_markdown(self):
        assert is_structured_markdown("Formula $x^2$ here") is True
