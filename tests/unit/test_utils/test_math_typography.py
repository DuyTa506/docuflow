"""Công thức fallback đang in bằng font thân bài, không có kiểu toán nào.

`_add_math_run` chỉ dựng OMML khi fragment có lệnh LaTeX thật (`\\Delta`); mọi
thứ còn lại — biến trần, mũ, chỉ số — đi qua `_add_latex_text_fallback`, vốn
`paragraph.add_run(...)` trần trụi. Kết quả: `n` in ra như chữ thường trong câu,
không phân biệt được với văn xuôi.

Quy ước sắp chữ toán (và cũng là thứ Word tự làm trong vùng OMML):
  - font **Cambria Math** — font toán chuẩn của Word, có đủ ký hiệu
  - **biến in nghiêng**, chữ số và toán tử đứng thẳng
  - **cỡ chữ kế thừa thân bài** — không đặt cứng, để công thức không to/nhỏ lệch
    so với câu chứa nó khi đổi style mẫu
"""

import pytest
from docx import Document

from utils.markdown_docx import _add_inline_runs

MATH_FONT = "Cambria Math"


def _runs(text: str):
    doc = Document()
    para = doc.add_paragraph()
    _add_inline_runs(para, text, strip_html=False)
    return para.runs


def _find(runs, text):
    return next((r for r in runs if (r.text or "") == text), None)


class TestVariableStyling:
    def test_variable_is_italic_in_the_math_font(self):
        run = _find(_runs("chương trình có $n$ dòng"), "n")

        assert run is not None, "biến phải thành run riêng"
        assert run.font.italic is True
        assert run.font.name == MATH_FONT

    def test_dollars_are_gone(self):
        assert "$" not in "".join(r.text or "" for r in _runs("có $n$ dòng"))

    def test_surrounding_prose_is_not_restyled(self):
        runs = _runs("chương trình có $n$ dòng")

        prose = [r for r in runs if "dòng" in (r.text or "")]
        assert prose and prose[0].font.italic is not True
        assert prose[0].font.name != MATH_FONT


class TestSuperSubscript:
    def test_exponent_letter_is_italic_math(self):
        run = _find(_runs("không gian $2^n$ ô nhớ"), "n")

        assert run is not None
        assert run.font.superscript is True
        assert run.font.italic is True
        assert run.font.name == MATH_FONT

    def test_digits_stay_upright(self):
        """`2` là hằng số — đứng thẳng, không nghiêng."""
        run = _find(_runs("không gian $2^n$ ô nhớ"), "2")

        assert run is not None
        assert run.font.italic is not True

    def test_subscript_keeps_math_font(self):
        run = _find(_runs("tần số $f_s$ lấy mẫu"), "s")

        assert run is not None
        assert run.font.subscript is True
        assert run.font.name == MATH_FONT


class TestSize:
    def test_size_is_inherited_not_pinned(self):
        """Đặt cứng cỡ chữ là sai: mẫu đổi style thì công thức lệch khỏi câu."""
        for r in _runs("có $n$ dòng và $2^n$ ô"):
            assert r.font.size is None


class TestCurrencyUntouched:
    def test_a_price_keeps_its_dollars_and_its_font(self):
        runs = _runs("giá khoảng $100 cho mỗi đơn vị")

        text = "".join(r.text or "" for r in runs)
        assert "$100" in text
        assert all(r.font.name != MATH_FONT for r in runs)
