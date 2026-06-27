"""Tests for structure-preserving DOCX export: grid tables, inline math, images.

These regression tests are grounded in real DeepSeek-OCR output behavior:
the model emits HTML tables with colspan/rowspan, literal ``\\n`` inside cells,
and inline ``$...$`` math fragments (including malformed ``^{{...}}``).
"""

import base64
from io import BytesIO

import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from PIL import Image

from utils.markdown_docx import _add_html_table, render_layout_elements_to_docx
from utils.math_omml import is_pandoc_available
from utils.translation_elements import TranslatedElementView


def _cell_span(cell):
    """Return (gridSpan, vMerge) parsed from a cell's tcPr XML."""
    tcpr = cell._tc.find(qn("w:tcPr"))
    grid_span, v_merge = 1, None
    if tcpr is not None:
        g = tcpr.find(qn("w:gridSpan"))
        if g is not None:
            grid_span = int(g.get(qn("w:val")) or 1)
        v = tcpr.find(qn("w:vMerge"))
        if v is not None:
            v_merge = v.get(qn("w:val")) or "continue"
    return grid_span, v_merge


def _tiny_png_b64() -> str:
    img = Image.new("RGB", (32, 16), (10, 120, 200))
    buf = BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode()


class TestGridTable:
    def test_colspan_and_rowspan_merge(self):
        html = (
            "<table style='margin: auto;'>"
            "<tr><td>STT</td><td>NOI DUNG</td><td>GIA</td><td>SL</td></tr>"
            "<tr><td rowspan='2'>1</td><td>Spec A</td><td>100</td><td>2</td></tr>"
            "<tr><td>Spec B</td><td>200</td><td>3</td></tr>"
            "<tr><td colspan='3'>TONG CONG</td><td>500</td></tr>"
            "</table>"
        )
        doc = DocxDocument()
        _add_html_table(doc, html)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 4
        assert len(table.columns) == 4

        # rowspan on "1" -> vertical merge present
        has_vmerge = any(
            _cell_span(table.cell(r, c))[1] is not None
            for r in range(4)
            for c in range(4)
        )
        # colspan on "TONG CONG" -> horizontal merge present
        has_gridspan = any(
            _cell_span(table.cell(r, c))[0] > 1
            for r in range(4)
            for c in range(4)
        )
        assert has_vmerge, "rowspan should produce a vertical merge"
        assert has_gridspan, "colspan should produce a horizontal merge"

        # the rowspan origin keeps the value '1'
        assert table.cell(1, 0).text.strip() == "1"
        # value column of totals row lands in the last column
        assert table.cell(3, 3).text.strip() == "500"

    def test_literal_backslash_n_becomes_linebreak(self):
        html = "<table><tr><td>DON GIA\\n(VND)</td><td>10-12\\nTuan</td></tr></table>"
        doc = DocxDocument()
        _add_html_table(doc, html)
        cell = doc.tables[0].cell(0, 0)
        assert "\n" in cell.text
        assert "\\n" not in cell.text

    def test_header_row_is_bold(self):
        html = "<table><tr><td>H1</td><td>H2</td></tr><tr><td>a</td><td>b</td></tr></table>"
        doc = DocxDocument()
        _add_html_table(doc, html)
        header_cell = doc.tables[0].cell(0, 0)
        assert any(run.bold for run in header_cell.paragraphs[0].runs)


class TestInlineMath:
    def test_malformed_superscript_renders_without_dollar(self):
        elems = [
            TranslatedElementView(
                label="text",
                text_content=r"5th Gen Intel $ ^{{®}} $ Xeon",
                page_number=1,
                bbox_x1=0, bbox_y1=0, bbox_x2=100, bbox_y2=10,
            )
        ]
        doc = DocxDocument()
        render_layout_elements_to_docx(doc, elems)
        full = "".join(p.text for p in doc.paragraphs)
        assert "$" not in full
        assert "®" in full
        sup = any(
            run.font.superscript
            for p in doc.paragraphs
            for run in p.runs
            if run.font.superscript
        )
        assert sup, "registered mark should be a superscript run"

    def test_currency_is_not_treated_as_math(self):
        elems = [
            TranslatedElementView(
                label="text",
                text_content="Gia $5 va $10 dong",
                page_number=1,
                bbox_x1=0, bbox_y1=0, bbox_x2=100, bbox_y2=10,
            )
        ]
        doc = DocxDocument()
        render_layout_elements_to_docx(doc, elems)
        full = "".join(p.text for p in doc.paragraphs)
        assert "$5" in full and "$10" in full

    def test_real_math_leaves_no_raw_dollar(self):
        elems = [
            TranslatedElementView(
                label="text",
                text_content=r"Dien tich $\frac{1}{2}bh$ la",
                page_number=1,
                bbox_x1=0, bbox_y1=0, bbox_x2=100, bbox_y2=10,
            )
        ]
        doc = DocxDocument()
        render_layout_elements_to_docx(doc, elems)
        para = doc.paragraphs[-1]
        assert "$" not in para.text
        if is_pandoc_available():
            omath = para._element.findall(qn("m:oMath")) + para._element.findall(qn("m:oMathPara"))
            assert len(omath) >= 1, "pandoc present: expected native OMML equation"


class TestImageEmbedding:
    def test_figure_embedded_as_picture_with_caption(self):
        b64 = _tiny_png_b64()
        elems = [
            TranslatedElementView(
                label="figure",
                text_content="Hinh 1: bieu do",
                page_number=1,
                bbox_x1=100, bbox_y1=50, bbox_x2=400, bbox_y2=250,
                crop_image_base64=b64,
            )
        ]
        doc = DocxDocument()
        render_layout_elements_to_docx(doc, elems)
        assert len(doc.inline_shapes) == 1
        assert any("Hinh 1" in p.text for p in doc.paragraphs)

    def test_image_placeholder_caption_suppressed(self):
        b64 = _tiny_png_b64()
        elems = [
            TranslatedElementView(
                label="image",
                text_content="(img_content)[image_1.png]",
                page_number=1,
                bbox_x1=100, bbox_y1=50, bbox_x2=400, bbox_y2=250,
                crop_image_base64=b64,
            )
        ]
        doc = DocxDocument()
        render_layout_elements_to_docx(doc, elems)
        assert len(doc.inline_shapes) == 1
        assert not any("img_content" in p.text for p in doc.paragraphs)

    def test_no_image_source_falls_back_to_text(self):
        elems = [
            TranslatedElementView(
                label="figure",
                text_content="Hinh khong co anh",
                page_number=1,
                bbox_x1=100, bbox_y1=50, bbox_x2=400, bbox_y2=250,
            )
        ]
        doc = DocxDocument()
        render_layout_elements_to_docx(doc, elems)
        assert len(doc.inline_shapes) == 0
        assert any("Hinh khong co anh" in p.text for p in doc.paragraphs)
