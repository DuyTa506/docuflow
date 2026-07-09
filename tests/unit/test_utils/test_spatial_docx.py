"""Tests for bbox-based spatial DOCX rendering."""

import io

from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from utils.markdown_docx import render_layout_elements_to_docx
from utils.translation_elements import TranslatedElementView


class TestSpatialDocx:
    def test_body_text_does_not_force_bbox_based_alignment(self):
        """Regression: exported DOCX should read as a standard document
        (Vietnamese convention: justified body text, first-line indent from
        the "Normal" style), not a geometric clone of the source PDF's bbox
        positions. Regular "text"-labeled elements must not get an explicit
        per-paragraph CENTER override just because they happened to sit
        near the horizontal center of the original page."""
        doc = DocxDocument()
        elements = [
            TranslatedElementView(
                label="text",
                text_content="Hebertt Sira-Ramírez",
                page_number=1,
                bbox_x1=200,
                bbox_y1=100,
                bbox_x2=400,
                bbox_y2=120,
            ),
            TranslatedElementView(
                label="text",
                text_content="Control Design",
                page_number=1,
                bbox_x1=180,
                bbox_y1=130,
                bbox_x2=420,
                bbox_y2=150,
            ),
        ]
        render_layout_elements_to_docx(doc, elements)
        body_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        assert body_paragraphs
        for p in body_paragraphs:
            # No explicit per-paragraph override -- alignment/indent come
            # from the "Normal" style default instead.
            assert p.alignment is None
        assert doc.styles["Normal"].paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def test_equation_element_renders(self):
        doc = DocxDocument()
        elements = [
            TranslatedElementView(
                label="formula",
                text_content="E = mc^2",
                page_number=1,
                bbox_x1=100,
                bbox_y1=50,
                bbox_x2=300,
                bbox_y2=70,
            ),
        ]
        render_layout_elements_to_docx(doc, elements)
        buf = io.BytesIO()
        doc.save(buf)
        assert len(buf.getvalue()) > 1000
