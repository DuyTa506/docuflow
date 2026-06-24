"""Tests for OCR markdown normalization and DOCX rendering."""

import io

from docx import Document as DocxDocument

from utils.file_download import build_docx_response
from utils.markdown_docx import build_docx_bytes_from_markdown, render_markdown_to_docx
from utils.ocr_markdown import is_structured_markdown, normalize_ocr_markdown


SAMPLE_OCR = """# QUYẾT ĐỊNH

## Điều 1. Phạm vi điều chỉnh

Quyết định này quy định về **nâng cao chất lượng** sinh hoạt chi bộ.

| STT | Nội dung | Ghi chú |
| --- | --- | --- |
| 1 | Tổ chức họp | Định kỳ |
| 2 | Báo cáo | Hàng quý |

<table><tr><td>A</td><td>B</td></tr><tr><td>1</td><td>2</td></tr></table>

- Mục một
- Mục hai
"""


class TestOcrMarkdown:
    def test_detects_structured_content(self):
        assert is_structured_markdown(SAMPLE_OCR) is True
        assert is_structured_markdown("Plain text only.\nSecond line.") is False

    def test_normalize_strips_grounding_tags(self):
        raw = "<|ref|>title<|/ref|><|det|>[[0,0,1,1]]<|/det|>\n# Title"
        cleaned = normalize_ocr_markdown(raw)
        assert "<|ref|>" not in cleaned
        assert "# Title" in cleaned


class TestMarkdownDocx:
    def test_renders_headings_and_tables(self):
        docx_bytes = build_docx_bytes_from_markdown(SAMPLE_OCR, title="Test Doc")
        doc = DocxDocument(io.BytesIO(docx_bytes))

        styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
        assert "Heading 1" in styles
        assert "Heading 2" in styles
        assert len(doc.tables) >= 2

        table_text = " ".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "STT" in table_text
        assert "A" in table_text
        assert "2" in table_text

    def test_page_breaks_on_separator(self):
        text = "Page one\n\n---\n\n# Page two"
        doc = DocxDocument()
        render_markdown_to_docx(doc, text)
        assert any(
            "w:br" in p._element.xml and 'type="page"' in p._element.xml for p in doc.paragraphs
        )

    def test_build_docx_response_uses_structured_path(self):
        resp = build_docx_response("ocr_test.docx", SAMPLE_OCR, title="Test")
        assert resp.media_type.endswith("wordprocessingml.document")
        doc = DocxDocument(io.BytesIO(resp.body))
        assert len(doc.tables) >= 1

    def test_plain_mode_for_unstructured_text(self):
        resp = build_docx_response(
            "plain.docx",
            "Line one\nLine two",
            structured=False,
        )
        doc = DocxDocument(io.BytesIO(resp.body))
        assert [p.text for p in doc.paragraphs] == ["Line one", "Line two"]
