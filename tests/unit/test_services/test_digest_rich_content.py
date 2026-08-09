"""Tổng thuật in ra `$\\Delta w$` nguyên chuỗi thay vì công thức.

Repo đã có sẵn đường xử lý toán — `utils/math_omml` dựng OMML qua pandoc, và
`utils/markdown_docx._add_inline_runs` tách `$...$` / `$$...$$` rồi chèn vào
đoạn Word. Bản tải văn bản dùng nó; **tổng thuật thì không**: `DigestRenderer`
chỉ gọi `plain_text()`, vốn gỡ markdown và để nguyên LaTeX.

Cách sửa: bắt tầng render Word đọc được thứ model sinh ra, thay vì bắt model
đừng sinh. Ràng buộc bằng prompt là thứ dễ vỡ nhất — đổi model một cái là hành
vi đổi theo (đo được khi chuyển Qwen → Gemma: Gemma viết LaTeX, Qwen không).

Ràng buộc phải giữ: mẫu quy định §2.2 là **một đoạn** `Chương N. Tên (Gốc).
Nội dung…` — nhãn và nội dung KHÔNG được tách thành hai đoạn.
"""

import io

import pytest
from docx import Document

from services.digest_renderer import DigestRenderer
from services.digest_service import (
    ChapterEntry,
    DigestResult,
    KeywordEntry,
    ResearchDirectionEntry,
)

_OMML = "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"


def _digest(content="Nội dung.", abstract="Tóm tắt."):
    return DigestResult(
        document_id="DOC_001",
        title="T",
        source_language="ru",
        original_filename=None,
        bibliographic={"title_display": "T", "pages": "1"},
        abstract=abstract,
        chapters=[
            ChapterEntry(
                number=1,
                title_vi="Giới thiệu",
                title_original="Введение",
                content=content,
                heading_kind="chapter",
                heading_ordinal=1,
            )
        ],
        keywords=[KeywordEntry(keyword="k", display="Từ khoá (k)", weight=0.9)],
        usage_scope={"undergraduate": [], "master": [], "phd": [], "strong_research_groups": []},
        research_directions=[ResearchDirectionEntry(direction_name="H", confidence=0.8)],
    )


def _doc(digest):
    return Document(io.BytesIO(DigestRenderer().render(digest)))


def _paras(digest):
    return [p.text for p in _doc(digest).paragraphs]


def _math_count(digest) -> int:
    return sum(1 for _ in _doc(digest).element.body.iter(_OMML))


class TestMath:
    def test_inline_formula_becomes_a_real_equation(self):
        digest = _digest(content=r"Đồng bộ qua khoảng $\Delta w$ trong một chu kỳ.")

        assert _math_count(digest) >= 1
        assert not any("$" in p for p in _paras(digest))

    def test_display_formula_becomes_a_real_equation(self):
        digest = _digest(content=r"Điều kiện lấy mẫu: $$f_s \geq 2 f_{max}$$ theo Nyquist.")

        assert _math_count(digest) >= 1

    def test_surrounding_prose_survives(self):
        digest = _digest(content=r"Đồng bộ qua khoảng $\Delta w$ trong một chu kỳ.")

        text = "\n".join(_paras(digest))
        assert "Đồng bộ qua khoảng" in text
        assert "trong một chu kỳ." in text

    def test_a_price_is_not_a_formula(self):
        """`looks_like_math` phải chặn — $100 là tiền, không phải toán."""
        digest = _digest(content="Chi phí khoảng $100 cho mỗi đơn vị.")

        assert _math_count(digest) == 0
        assert any("$100" in p for p in _paras(digest))

    def test_abstract_gets_the_same_treatment(self):
        digest = _digest(abstract=r"Tài liệu dùng ký hiệu $\Delta t$ làm ví dụ.")

        assert _math_count(digest) >= 1

    def test_simple_exponent_becomes_a_superscript_not_an_equation(self):
        """`_add_math_run` chỉ gọi pandoc khi có lệnh LaTeX thật; `2^n` thì dựng
        run chỉ số trên. Rẻ hơn và đọc ra vẫn đúng — cái phải biến mất là `$`."""
        digest = _digest(content=r"Không gian địa chỉ $2^n$ ô nhớ.")

        doc = _doc(digest)
        assert not any("$" in p.text for p in doc.paragraphs)
        assert any(
            r.font.superscript and (r.text or "").strip() == "n"
            for p in doc.paragraphs
            for r in p.runs
        )


class TestMarkdownNowRendered:
    def test_bold_becomes_bold_not_literal_and_not_stripped(self):
        # Bold sits mid-sentence so this stays a test about bold alone —
        # a word in the opening position is also subject to the opener rule.
        digest = _digest(content="Trình bày phần **nhấn mạnh** kiến trúc.")

        doc = _doc(digest)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "**" not in text.replace("*****", "")  # masthead có ***** thật
        assert "nhấn mạnh" in text
        assert any(r.bold and "nhấn mạnh" in (r.text or "") for p in doc.paragraphs for r in p.runs)

    def test_block_heading_markers_are_still_removed(self):
        digest = _digest(abstract="## Chủ đề chính\n\nNội dung.")

        paras = _paras(digest)
        assert "Chủ đề chính" in paras
        assert not any(p.startswith("#") for p in paras)


class TestTemplateShapeUnchanged:
    def test_heading_and_content_stay_in_one_paragraph(self):
        """Mẫu: `Chương 1. Giới thiệu (Введение). Nội dung…` — MỘT đoạn."""
        digest = _digest(content="Nội dung chương.")

        assert "Chương 1. Giới thiệu (Введение). Nội dung chương." in _paras(digest)

    def test_no_blank_paragraph_inflation(self):
        small = _paras(_digest())
        big = _digest()
        big.chapters = big.chapters * 6
        for i, c in enumerate(big.chapters, 1):
            c.number = i
            c.heading_ordinal = i

        blanks = lambda ps: sum(1 for p in ps if not p.strip())  # noqa: E731
        assert blanks(small) == blanks(_paras(big))

    def test_missing_abstract_still_says_so(self):
        digest = _digest()
        digest.abstract = None

        assert "[Chưa có — chạy summarize trước]" in "\n".join(_paras(digest))
