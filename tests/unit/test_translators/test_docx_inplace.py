"""Tests for DocxInPlaceTranslator."""

import io
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document as DocxDocument

from services.translators.docx_inplace_translator import DocxInPlaceTranslator


@pytest.fixture
def sample_docx(tmp_path):
    path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_heading("Hello", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Cell A"
    table.rows[0].cells[1].text = "Cell B"
    doc.save(path)
    return path


class TestDocxInPlaceTranslator:
    @pytest.mark.asyncio
    async def test_translates_paragraphs_and_tables(self, sample_docx, tmp_path):
        out = tmp_path / "out.docx"
        translator = MagicMock()
        translator.translate_text = AsyncMock(side_effect=lambda t: f"VI:{t}")

        svc = DocxInPlaceTranslator(translator)
        result = await svc.translate_file(str(sample_docx), str(out), doc_format="docx")

        assert result["translation_mode"] == "docx_inplace"
        assert out.exists()

        doc = DocxDocument(str(out))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("VI:Hello" in t for t in texts)

        table_text = " ".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "VI:Cell A" in table_text
        assert "VI:Cell B" in table_text


class TestParagraphConcurrency:
    @pytest.mark.asyncio
    async def test_paragraphs_translate_concurrently_and_keep_document_order(self, tmp_path):
        """Paragraphs were translated strictly one-at-a-time — a multi-hundred
        page DOCX must fan out (bounded) while output order stays exact."""
        import asyncio

        path = tmp_path / "many.docx"
        doc = DocxDocument()
        for i in range(8):
            doc.add_paragraph(f"Paragraph number {i}")
        doc.save(path)
        out = tmp_path / "out.docx"

        in_flight = [0]
        peak = [0]

        async def fake_translate(text):
            in_flight[0] += 1
            peak[0] = max(peak[0], in_flight[0])
            await asyncio.sleep(0.02)
            in_flight[0] -= 1
            return f"VI:{text}"

        translator = MagicMock()
        translator.translate_text = AsyncMock(side_effect=fake_translate)

        svc = DocxInPlaceTranslator(translator)
        result = await svc.translate_file(str(path), str(out), doc_format="docx")

        assert peak[0] > 1, "paragraphs still translate one at a time"
        texts = [p.text for p in DocxDocument(str(out)).paragraphs if p.text.strip()]
        assert texts == [f"VI:Paragraph number {i}" for i in range(8)]
        flat = result["translated_content"].split("\n\n")
        assert flat == [f"VI:Paragraph number {i}" for i in range(8)]


class TestRunFormattingPreservation:
    """_set_para_text used to dump the whole translation into run 0 and blank
    every other run — destroying bold/italic/hyperlink formatting inside the
    paragraph."""

    def _para_with_bold_middle(self):
        doc = DocxDocument()
        p = doc.add_paragraph()
        p.add_run("The quick brown ")
        bold = p.add_run("fox jumps")
        bold.bold = True
        p.add_run(" over the lazy dog")
        return doc, p

    def test_mixed_format_paragraph_keeps_all_runs_texted(self):
        doc, p = self._para_with_bold_middle()
        translated = "Con cáo nâu nhanh nhẹn nhảy qua con chó lười biếng nằm im"

        DocxInPlaceTranslator._set_para_text(p._p, translated)

        run_texts = [r.text for r in p.runs]
        assert all(t.strip() for t in run_texts), f"blanked runs: {run_texts}"
        # bold formatting stays on the middle run
        assert p.runs[1].bold is True
        # concatenation reads as the full translation
        assert p.text.split() == translated.split()

    def test_single_run_paragraph_unchanged_behavior(self):
        doc = DocxDocument()
        p = doc.add_paragraph()
        p.add_run("Only one run here")

        DocxInPlaceTranslator._set_para_text(p._p, "Chỉ một run")

        assert p.text == "Chỉ một run"

    def test_hyperlink_run_not_blanked(self):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        doc = DocxDocument()
        p = doc.add_paragraph()
        p.add_run("See the ")
        hyperlink = OxmlElement("w:hyperlink")
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "documentation page"
        run.append(t)
        hyperlink.append(run)
        p._p.append(hyperlink)
        p.add_run(" for details")

        DocxInPlaceTranslator._set_para_text(
            p._p, "Xem trang tài liệu hướng dẫn để biết thêm chi tiết"
        )

        link_texts = [
            el.text or ""
            for el in p._p.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
                "//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
            )
        ]
        assert link_texts and all(
            t.strip() for t in link_texts
        ), f"hyperlink text blanked: {link_texts}"
